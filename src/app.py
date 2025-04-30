from flask import Flask, render_template, request, send_file
from io import BytesIO
import weasyprint
from docx import Document
from docx.shared import Pt

app = Flask(__name__, template_folder="templates")

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Form data
        data = {
            "name": request.form.get("name", ""),
            "email": request.form.get("email", ""),
            "phone": request.form.get("phone", ""),
            "address": request.form.get("address", ""),
            "website": request.form.get("website", ""),
            "dob": request.form.get("dob", ""),
            "gender": request.form.get("gender", ""),
            "summary": request.form.get("summary", ""),
            "experience": request.form.get("experience", ""),
            "education": request.form.get("education", ""),
            "skills": request.form.get("skills", ""),
            "languages": request.form.get("languages", ""),
            "certifications": request.form.get("certifications", "")
        }

        # Get template and format from the form data
        chosen_template = request.form.get("template", "cv_output")
        format_type = request.form.get("format", "pdf")  # Get format from POST data

        # Generate PDF
        if format_type == "pdf":
            html = render_template(f"{chosen_template}.html", data=data)
            pdf = weasyprint.HTML(string=html).write_pdf()
            pdf_io = BytesIO(pdf)
            pdf_io.seek(0)
            return send_file(pdf_io, as_attachment=True, download_name="cv_output.pdf", mimetype="application/pdf")

        # Generate DOCX
        elif format_type == "docx":
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            doc.add_heading(data["name"], 0)
            doc.add_paragraph(f"Email: {data['email']}")
            doc.add_paragraph(f"Phone: {data['phone']}")
            doc.add_paragraph(f"Address: {data['address']}")
            doc.add_paragraph(f"Website: {data['website']}")
            doc.add_paragraph(f"DOB: {data['dob']} | Gender: {data['gender']}")

            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data["summary"])

            doc.add_heading("Experience", level=1)
            doc.add_paragraph(data["experience"])

            doc.add_heading("Education", level=1)
            doc.add_paragraph(data["education"])

            doc.add_heading("Skills", level=1)
            doc.add_paragraph(data["skills"])

            doc.add_heading("Languages", level=1)
            doc.add_paragraph(data["languages"])

            doc.add_heading("Certifications", level=1)
            for cert in data["certifications"].split("\n"):
                doc.add_paragraph(cert.strip(), style='List Bullet')

            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return send_file(
                doc_io,
                as_attachment=True,
                download_name="cv_output.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

if __name__ == "__main__":
    app.run(debug=False)
